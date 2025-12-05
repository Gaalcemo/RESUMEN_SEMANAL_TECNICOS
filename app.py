import streamlit as st
import pandas as pd
import datetime as dt
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional
import re
import calendar
import io

import gspread
from google.oauth2.service_account import Credentials
import holidays

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# === CONFIGURACIÓN BÁSICA ===

SPREADSHEETS = {
    2025: "1x2m_seSTeY9O7mE_ZqHJy36KQ2dvYmlvtIjCXLe1TyM",
    2026: "1MdS9nLp3gPSeWXbkEQKsffNyleIyGOytu7ySJ69xMtY",
}

MONTH_CODES = {
    1: "ENE",
    2: "FEB",
    3: "MAR",
    4: "ABR",
    5: "MAY",
    6: "JUN",
    7: "JUL",
    8: "AGO",
    9: "SEP",
    10: "OCT",
    11: "NOV",
    12: "DIC",
}

WEEKDAY_LETTER = {0: "L", 1: "M", 2: "X", 3: "J", 4: "V", 5: "S", 6: "D"}


# === SIGLAS Y AGRUPACIONES ===

SIGLA_DESCRIPTIONS: Dict[str, str] = {
    "L31": "cocacola reparacion",
    "MV": "movitech",
    "RY": "reyenvas",
    "/": "",
    "AC": "angel camacho",
    "VR": "vacaciones",
    "AR": "aranco",
    "BF": "bidafarma",
    "BFp": "bidafarma puerto real",
    "BFh": "bidafarma huelva",
    "BFm": "bidafarma malaga",
    "BFcm": "bidafarma camas",
    "BOR": "bordas",
    "HISP": "hispacold",
    "AF": "alfran",
    "AL": "alliance",
    "LAMt": "la muralla",
    "LAMm": "la muralla",
    "FOR": "formacion",
    "LYS": "lysur",
    "HA": "hariberica",
    "ST": "strugal",
    "CEU": "ceu",
    "MC": "mecalux",
    "SAI": "saica",
    "VEOt": "veolia",
    "BP": "baja paternidad",
    "PER": "persan fijos",
    "PERt": "persan turnos",
    "PERm": "persan turnos",
    "PERn": "persan turnos",
    "FORM": "formacion",
    "D": "libre",
    "CCs": "cocacola sevilla",
    "BE": "baja enfermedad",
    "BE ?": "baja enfermedad por confirmar",
    "CCsm": "cocacola sevilla",
    "P": "refresco iberia",
    "YB": "ybarra",
    "PRM": "primor",
    "LD": "lidl",
    "ONET": "onet",
    "VEI": "veimancha",
    "TH": "healthcare",
    "AP": "asuntos propios",
    "F5": "factor 5",
    "MED": "reconocimiento medico",
    "SMm": "sanmiguel",
    "NAVE": "nave",
    "L10Ele": "cocacola reparacion",
    "IN": "inalcoa",
    "ITP": "itp",
    "CP": "cementos portland",
    "M1": "persan turnos",
    "T1": "persan turnos",
    "N1": "persan turnos",
    "M2": "persan turnos",
    "T2": "persan turnos",
    "N2": "persan turnos",
    "M3": "persan turnos",
    "T3": "persan turnos",
    "N3": "persan turnos",
    "M4": "persan turnos",
    "T4": "persan turnos",
    "N4": "persan turnos",
    "ITV": "itv",
}

SIGLA_COMPANY: Dict[str, str] = {
    "PER": "Persan fijos",
    "PERt": "Persan turnos",
    "PERm": "Persan turnos",
    "PERn": "Persan turnos",
}

SEVILLA_LOCAL_HOLIDAYS: Dict[dt.date, str] = {
    # dt.date(2025, 4, 23): "Fiesta local Sevilla",
}


# === ESTRUCTURAS DE DATOS ===

@dataclass
class MonthSheet:
    df: pd.DataFrame
    day_col_map: Dict[int, int]
    tech_row_map: Dict[str, int]
    tech_names: List[str]


# === FUNCIONES AUXILIARES ===

def normalize_name(name: str) -> str:
    return re.sub(r"\s+", " ", name.strip()).upper()


def is_person_name(name: str) -> bool:
    n = name.strip()
    if not n:
        return False
    if re.search(r"\d", n):
        return False
    parts = [p for p in re.split(r"\s+", n) if p]
    if len(parts) < 2:
        return False
    alpha_tokens = sum(
        1 for p in parts if re.search(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]", p)
    )
    return alpha_tokens >= 2


@st.cache_resource
def get_gspread_client():
    import json
    from google.oauth2.service_account import Credentials

    scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]

    # Leemos el bloque [gcp_service_account] de secrets
    service_info = dict(st.secrets["gcp_service_account"])

    creds = Credentials.from_service_account_info(service_info, scopes=scopes)
    client = gspread.authorize(creds)
    return client

def find_day_row(df: pd.DataFrame) -> Optional[int]:
    for idx in range(min(10, len(df))):
        row = df.iloc[idx]
        count_numeric = 0
        for cell in row[2:]:
            cell_str = str(cell).strip()
            if cell_str.isdigit():
                count_numeric += 1
        if count_numeric >= 5:
            return idx
    return None


@st.cache_data
def load_month_sheet(year: int, month: int) -> MonthSheet:
    client = get_gspread_client()

    if year not in SPREADSHEETS:
        raise ValueError(f"No hay configuración para el año {year}")

    spreadsheet_id = SPREADSHEETS[year]
    month_code = MONTH_CODES[month]

    sh = client.open_by_key(spreadsheet_id)
    ws = sh.worksheet(month_code)

    data = ws.get_all_values()
    df = pd.DataFrame(data)

    day_row_idx = find_day_row(df)
    if day_row_idx is None:
        raise ValueError("No se encontró la fila de días (1..31) en la hoja.")

    tech_start = day_row_idx + 2

    day_col_map: Dict[int, int] = {}
    header_row = df.iloc[day_row_idx]
    for col_idx in range(2, len(header_row)):
        cell_str = str(header_row.iloc[col_idx]).strip()
        if cell_str.isdigit():
            day = int(cell_str)
            day_col_map[day] = col_idx

    tech_row_map: Dict[str, int] = {}
    tech_names: List[str] = []

    last_row_to_read = min(len(df), 82)

    for row_idx in range(tech_start, last_row_to_read):
        raw_name = str(df.iat[row_idx, 1]).strip()
        if not raw_name:
            continue
        if not is_person_name(raw_name):
            continue

        tech_names.append(raw_name)
        tech_row_map[normalize_name(raw_name)] = row_idx

    return MonthSheet(df=df, day_col_map=day_col_map, tech_row_map=tech_row_map, tech_names=tech_names)


def get_assignment(month_sheet: MonthSheet, tech_name: str, day: int) -> str:
    tech_row = month_sheet.tech_row_map.get(normalize_name(tech_name))
    col_idx = month_sheet.day_col_map.get(day)
    if tech_row is None or col_idx is None:
        return ""
    value = str(month_sheet.df.iat[tech_row, col_idx]).strip()
    if value == "D":
        return ""
    return value


def describe_sigla(code: str) -> str:
    if not code:
        return ""
    code = code.strip()
    desc = SIGLA_DESCRIPTIONS.get(code)
    if desc is not None:
        return desc

    m_l = re.match(r"^L(\d+)", code)
    if m_l:
        num = int(m_l.group(1))
        if num < 50:
            return "cocacola reparacion"
        elif num > 51:
            return "refresco reparacion"

    if re.match(r"^[MTN]\d+$", code):
        return "persan turnos"

    return code


def get_company_for_code(code: str) -> str:
    if not code:
        return ""
    code = code.strip()

    company = SIGLA_COMPANY.get(code)
    if company:
        return company

    if re.match(r"^[MTN]\d+$", code):
        return "Persan turnos"

    m_l = re.match(r"^L(\d+)", code)
    if m_l:
        num = int(m_l.group(1))
        if num < 50:
            return "Cocacola reparacion"
        elif num > 51:
            return "Refresco reparacion"

    desc = SIGLA_DESCRIPTIONS.get(code)
    if desc:
        return desc.capitalize()

    return code


def get_holidays_for_dates(dates: List[dt.date]) -> Dict[dt.date, str]:
    years = sorted({d.year for d in dates})
    es_holidays = holidays.Spain(years=years, subdiv="AN")
    festivos: Dict[dt.date, str] = {}
    for d in dates:
        if d in es_holidays:
            festivos[d] = str(es_holidays.get(d))
    for d, desc in SEVILLA_LOCAL_HOLIDAYS.items():
        if d in dates:
            festivos[d] = desc
    return festivos


def compress_days(indices: List[int]) -> str:
    unique = sorted(set(indices))
    if not unique:
        return ""
    min_i, max_i = unique[0], unique[-1]
    if len(unique) == (max_i - min_i + 1) and len(unique) > 1:
        return f"{WEEKDAY_LETTER[min_i]} - {WEEKDAY_LETTER[max_i]}"
    elif len(unique) == 1:
        return WEEKDAY_LETTER[unique[0]]
    else:
        return ",".join(WEEKDAY_LETTER[i] for i in unique)


def build_resumen_lines(daily_info: List[Tuple[dt.date, str, str, bool]]) -> List[str]:
    if not daily_info:
        return []

    company_days: Dict[str, Dict] = {}

    for d, label, desc, _is_fest in daily_info:
        if label == "LIBRE":
            company_key = "LIBRE"
            company_display = "LIBRE"
        else:
            company_display = get_company_for_code(label)
            company_key = company_display.upper()

        wd = d.weekday()
        entry = company_days.setdefault(
            company_key,
            {"company_display": company_display, "days": [], "codes": set()},
        )
        entry["days"].append((wd, label))
        if label != "LIBRE":
            entry["codes"].add(label)

    ordered_groups = sorted(
        company_days.values(),
        key=lambda e: min(wd for wd, _ in e["days"])
    )

    lines: List[str] = []

    for group in ordered_groups:
        company_display = group["company_display"]
        days = group["days"]
        codes = group["codes"]

        indices = [wd for wd, _ in days]
        days_str = compress_days(indices)

        if company_display == "LIBRE":
            label_text = "LIBRE"
        else:
            if len(codes) == 1:
                code = next(iter(codes))
                label_text = f"{code} ({company_display})"
            else:
                label_text = f"({company_display})"

        lines.append(f"{days_str} : {label_text}")

    return lines


def get_main_company_from_daily_info(daily_info: List[Tuple[dt.date, str, str, bool]]) -> str:
    companies: List[str] = []
    for _d, label, _desc, _fest in daily_info:
        if label == "LIBRE" or not label:
            continue
        companies.append(get_company_for_code(label))

    if not companies:
        return "LIBRE"

    counts: Dict[str, int] = {}
    for c in companies:
        counts[c] = counts.get(c, 0) + 1

    main_company = sorted(counts.items(), key=lambda x: (-x[1], x[0]))[0][0]
    return main_company


def get_weeks_in_month(year: int, month: int) -> List[Tuple[dt.date, dt.date]]:
    last_day = calendar.monthrange(year, month)[1]
    weeks: List[Tuple[dt.date, dt.date]] = []
    day = 1
    while day <= last_day:
        start = dt.date(year, month, day)
        end_day = min(day + 6, last_day)
        end = dt.date(year, month, end_day)
        weeks.append((start, end))
        day += 7
    return weeks


# === GENERACIÓN DE WORD (lista en columnas, empresa en negrita) ===

def build_word_doc(
    start_date: dt.date,
    end_date: dt.date,
    companies_order: List[str],
    selected_tecs: List[str],
    resumen_por_tecnico: Dict[str, List[str]],
    main_company_by_tech: Dict[str, str],
) -> bytes:
    doc = Document()

    # A4 vertical
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    # 2 columnas
    sectPr = section._sectPr
    cols_elems = sectPr.xpath("./w:cols")
    if cols_elems:
        cols = cols_elems[0]
    else:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")

    # Fuente general
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(8)

    # Título
    title = doc.add_paragraph()
    run = title.add_run(
        f"Resumen semanal {start_date.strftime('%d/%m/%Y')} - {end_date.strftime('%d/%m/%Y')}"
    )
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    # Contenido
    for company in companies_order:
        techs_in_company = sorted(
            [t for t in selected_tecs if main_company_by_tech[t] == company]
        )
        if not techs_in_company:
            continue

        # Empresa en negrita
        p_company = doc.add_paragraph()
        p_company.paragraph_format.space_before = Pt(4)
        p_company.paragraph_format.space_after = Pt(2)
        r_company = p_company.add_run(company)
        r_company.bold = True

        # Técnicos debajo (sin negrita, poco espacio entre líneas)
        for tech in techs_in_company:
            lines = resumen_por_tecnico.get(tech, [])
            if not lines:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p_run = p.add_run(f"{tech}: " + " ; ".join(lines))

        # Espacio entre empresas
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# === INTERFAZ STREAMLIT ===

def main():
    st.set_page_config(page_title="Resumen semanal técnicos", layout="wide")

    st.title("📅 Resumen semanal de asignaciones de técnicos")

    st.sidebar.header("Configuración")

    years = sorted(SPREADSHEETS.keys())
    year = st.sidebar.selectbox("Año", options=years, index=0)

    month = st.sidebar.selectbox(
        "Mes",
        options=list(range(1, 13)),
        format_func=lambda m: MONTH_CODES[m],
    )

    month_weeks = get_weeks_in_month(year, month)
    week_options = list(range(len(month_weeks)))
    week_labels = [
        f"Semana {i+1}: {start.day:02d}-{end.day:02d}"
        for i, (start, end) in enumerate(month_weeks)
    ]
    week_index = st.sidebar.selectbox(
        "Semana del mes",
        options=week_options,
        format_func=lambda i: week_labels[i],
    )

    start_date, end_date = month_weeks[week_index]
    week_dates = [
        start_date + dt.timedelta(days=i)
        for i in range((end_date - start_date).days + 1)
    ]

    st.markdown(
        f"**Semana seleccionada:** "
        f"{start_date.strftime('%d/%m/%Y')} → {end_date.strftime('%d/%m/%Y')}"
    )

    try:
        month_sheet = load_month_sheet(year, month)
    except Exception as e:
        st.error(f"No se pudo cargar la hoja para {MONTH_CODES[month]} {year}: {e}")
        return

    all_techs = sorted(month_sheet.tech_names)

    selected_tecs = st.sidebar.multiselect(
        "Técnicos a mostrar",
        options=all_techs,
        default=all_techs,
    )

    if not selected_tecs:
        st.info("Selecciona al menos un técnico en la barra lateral.")
        return

    festivos = get_holidays_for_dates(week_dates)

    resumen_por_tecnico: Dict[str, List[str]] = {}
    main_company_by_tech: Dict[str, str] = {}

    for tech in selected_tecs:
        daily_info_all: List[Tuple[dt.date, str, str, bool]] = []
        daily_info_resumen: List[Tuple[dt.date, str, str, bool]] = []

        for d in week_dates:
            assignment = get_assignment(month_sheet, tech, d.day)

            if not assignment:
                label = "LIBRE"
                desc = "LIBRE"
            else:
                label = assignment
                desc = describe_sigla(assignment)

            is_fest = d in festivos
            daily_info_all.append((d, label, desc, is_fest))

            if ((d.weekday() in (5, 6)) or is_fest) and label == "LIBRE":
                continue

            daily_info_resumen.append((d, label, desc, is_fest))

        resumen_por_tecnico[tech] = build_resumen_lines(daily_info_resumen)
        main_company_by_tech[tech] = get_main_company_from_daily_info(daily_info_all)

    companies_order = sorted({main_company_by_tech[t] for t in selected_tecs})

    # TXT (opcional)
    download_lines: List[str] = []
    header = (
        f"Resumen semanal {start_date.strftime('%d/%m/%Y')} - "
        f"{end_date.strftime('%d/%m/%Y')}\n"
    )
    download_lines.append(header)

    for company in companies_order:
        techs_in_company = sorted(
            [t for t in selected_tecs if main_company_by_tech[t] == company]
        )
        if not techs_in_company:
            continue
        download_lines.append(company.upper())
        for tech in techs_in_company:
            lines = resumen_por_tecnico.get(tech, [])
            if not lines:
                continue
            download_lines.append(f"{tech}")
            for line in lines:
                download_lines.append(f"  {line}")
            download_lines.append("")
        download_lines.append("")

    download_text = "\n".join(download_lines)
    st.download_button(
        "⬇️ Descargar resumen (.txt)",
        data=download_text,
        file_name="resumen_semana_tecnicos.txt",
        mime="text/plain",
    )

    # WORD
    word_bytes = build_word_doc(
        start_date,
        end_date,
        companies_order,
        selected_tecs,
        resumen_por_tecnico,
        main_company_by_tech,
    )

    st.download_button(
        "⬇️ Descargar resumen en Word (A4 columnas)",
        data=word_bytes,
        file_name="resumen_semana_tecnicos.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    # Vista en pantalla
    st.subheader("📊 Resumen agrupado por empresa y técnico")

    num_cols = 3
    for company in companies_order:
        techs_in_company = sorted(
            [t for t in selected_tecs if main_company_by_tech[t] == company]
        )
        if not techs_in_company:
            continue

        st.markdown(f"### {company}")

        cols = st.columns(num_cols)
        idx = 0

        for tech in techs_in_company:
            lines = resumen_por_tecnico.get(tech, [])
            col = cols[idx % num_cols]
            idx += 1

            with col:
                if not lines:
                    content_html = "Sin datos para esta semana."
                else:
                    content_html = "<br>".join(lines)

                st.markdown(
                    f"""
                    <div style="
                        border: 2px solid #333;
                        border-radius: 12px;
                        padding: 10px 14px;
                        margin: 8px 0;
                        min-height: 80px;
                        font-size: 0.9rem;
                        ">
                        <div style="font-weight: 700; margin-bottom: 6px;">
                            {tech}
                        </div>
                        <div style="line-height: 1.35;">
                            {content_html}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )


if __name__ == "__main__":
    main()

